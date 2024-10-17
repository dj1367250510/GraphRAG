import fitz
from paddleocr import PaddleOCR
from PIL import Image
import io
from docx import Document
import logging
from tqdm import tqdm

# 设置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# 初始化 PaddleOCR，选择中文模型 'ch' 表示中英文模型
ocr = PaddleOCR(use_angle_cls=True, lang='ch')

# 渲染 PDF 页面为图像并使用 PaddleOCR 进行文本识别
def render_page_and_ocr_paddle(pdf_path, output_docx_path):
    try:
        pdf_document = fitz.open(pdf_path)
    except Exception as e:
        logging.error(f"无法打开 PDF 文件: {e}")
        return

    doc = Document()
    doc.add_heading('页面 OCR 提取部分 (PaddleOCR)', level=1)

    # 使用 tqdm 添加进度条显示
    for page_num in tqdm(range(pdf_document.page_count), desc="正在处理 PDF 页面"):
        try:
            page = pdf_document[page_num]
            zoom = 2
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            image = Image.open(io.BytesIO(pix.tobytes()))

            # 将图像转换为可以被 PaddleOCR 处理的格式（保存为文件或处理字节流）
            image_byte_arr = io.BytesIO()
            image.save(image_byte_arr, format='PNG')
            image_byte_arr = image_byte_arr.getvalue()

            # 使用 PaddleOCR 进行文本识别
            result = ocr.ocr(image_byte_arr)

            # 提取并保存识别结果
            doc.add_heading(f'Page {page_num + 1} OCR Text from Rendered Image', level=2)
            for line in result:
                text = ''.join([word_info[1][0] for word_info in line])
                doc.add_paragraph(text)

            # 释放图像内存
            image.close()

        except Exception as e:
            logging.error(f"处理第 {page_num + 1} 页时出错: {e}")
            continue

    pdf_document.close()

    # 保存 Word 文件
    try:
        doc.save(output_docx_path)
        logging.info(f"页面 OCR 提取完成并保存至 {output_docx_path}")
    except Exception as e:
        logging.error(f"保存 Word 文件时出错: {e}")

# 调用主函数
pdf_path = r'E:\PythonCode\Graph RGA\pdf\test.pdf'
output_docx_path = r'E:\PythonCode\Graph RGA\word\output_document_paddleocr.docx'
render_page_and_ocr_paddle(pdf_path, output_docx_path)
