import fitz
import pytesseract  # OCR 引擎
from PIL import Image
import io
from docx import Document

# 提取 PDF 文本并保存到 Word 文档
def extract_text_from_pdf(pdf_path, output_docx_path):
    pdf_document = fitz.open(pdf_path)
    doc = Document()
    doc.add_heading('PDF 文本提取部分', level=1)

    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        text = page.get_text("text")  # 默认纯文本
        if text.strip() == "":
            text = page.get_text("html")  # 尝试使用 HTML 模式
            if text.strip() == "":
                text = page.get_text("dict")  # 尝试使用字典模式

        doc.add_heading(f'Page {page_num + 1}', level=2)
        doc.add_paragraph(text)

    pdf_document.close()
    doc.save(output_docx_path)
    print(f"PDF 文本提取完成并保存至 {output_docx_path}")

# 提取 PDF 图片并使用 OCR 进行文本识别，保存到 Word 文档
def extract_images_with_ocr(pdf_path, output_docx_path):
    pdf_document = fitz.open(pdf_path)
    doc = Document(output_docx_path)

    doc.add_heading('图片 OCR 提取部分', level=1)

    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        images = page.get_images(full=True)

        for img_index, img in enumerate(images):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(io.BytesIO(image_bytes))

            # 提高图像质量后使用多语言 OCR，识别英文和中文
            ocr_text = pytesseract.image_to_string(image, lang='eng+chi_sim')
            doc.add_heading(f'Page {page_num + 1}, Image {img_index + 1} OCR Text', level=2)
            doc.add_paragraph(ocr_text)

    pdf_document.close()
    doc.save(output_docx_path)
    print(f"图片 OCR 提取完成并保存至 {output_docx_path}")

# 将 PDF 页面渲染为图像并通过 OCR 处理，保存到 Word 文档
def render_page_as_image_and_ocr(pdf_path, output_docx_path):
    pdf_document = fitz.open(pdf_path)
    doc = Document(output_docx_path)

    doc.add_heading('页面 OCR 提取部分', level=1)

    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        zoom = 6
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        image = Image.open(io.BytesIO(pix.tobytes()))

        # 提高图像质量，使用 OCR 提取渲染的图像中的文本
        ocr_text = pytesseract.image_to_string(image, lang='chi_sim+eng')
        doc.add_heading(f'Page {page_num + 1} OCR Text from Rendered Image', level=2)
        doc.add_paragraph(ocr_text)

    pdf_document.close()
    doc.save(output_docx_path)
    print(f"页面 OCR 提取完成并保存至 {output_docx_path}")
