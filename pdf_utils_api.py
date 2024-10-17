from fastapi import FastAPI, UploadFile, File
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
import os
import fitz  # PyMuPDF
from docx import Document
import signal
import sys
from datetime import datetime

app = FastAPI()

# 信号处理，捕获 Ctrl+C
def handle_sigint(signal_received, frame):
    print("SIGINT or CTRL-C detected. Exiting gracefully...")
    sys.exit(0)

signal.signal(signal.SIGINT, handle_sigint)

# 将 'index.html' 文件所在的目录挂载为静态文件目录
app.mount("/static", StaticFiles(directory="static"), name="static")

# 定义文件保存的文件夹路径为 'docus'
UPLOAD_FOLDER = "docus/"
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# 根目录返回 HTML 文件
@app.get("/", response_class=FileResponse)
async def serve_homepage():
    return FileResponse("static/index.html")

# 提取 PDF 文本并保存到 Word 文档
def extract_text_from_pdf(pdf_path, output_docx_path):
    pdf_document = fitz.open(pdf_path)
    doc = Document()
    doc.add_heading('PDF 文本提取部分', level=1)

    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        text = page.get_text("text") or page.get_text("html") or page.get_text("dict")
        doc.add_heading(f'Page {page_num + 1}', level=2)
        doc.add_paragraph(text)

    pdf_document.close()
    doc.save(output_docx_path)

# 上传并保存文件到 'docus' 文件夹，并返回下载链接的 JSON
@app.post("/upload/", response_class=JSONResponse)
async def upload_pdf(file: UploadFile = File(...)):
    # 获取当前时间戳并生成唯一文件名
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    original_filename = file.filename
    base_name, ext = os.path.splitext(original_filename)

    # 定义带时间戳的文件名
    pdf_filename = f"{base_name}_{timestamp}{ext}"
    pdf_path = os.path.join(UPLOAD_FOLDER, pdf_filename)
    word_filename = f"{base_name}_{timestamp}.docx"
    output_docx_path = os.path.join(UPLOAD_FOLDER, word_filename)

    # 保存上传的 PDF 文件
    with open(pdf_path, "wb") as f:
        f.write(await file.read())

    # 处理 PDF 文件，生成 Word 文档
    extract_text_from_pdf(pdf_path, output_docx_path)

    # 返回 JSON 响应，包含下载链接
    download_word_link = f"/download/word/{word_filename}"
    download_pdf_link = f"/download/pdf/{pdf_filename}"

    return JSONResponse({
        "message": "File uploaded and processed successfully!",
        "download_word_link": download_word_link,
        "download_pdf_link": download_pdf_link
    })

# 提供下载生成的 Word 文件
@app.get("/download/word/{file_name}")
async def download_word_file(file_name: str):
    file_path = os.path.join(UPLOAD_FOLDER, file_name)
    if os.path.exists(file_path):
        return FileResponse(file_path,
                            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            filename=file_name)
    else:
        return JSONResponse({"message": "File not found"}, status_code=404)

# 提供下载原始 PDF 文件
@app.get("/download/pdf/{file_name}")
async def download_pdf_file(file_name: str):
    file_path = os.path.join(UPLOAD_FOLDER, file_name)
    if os.path.exists(file_path):
        return FileResponse(file_path,
                            media_type="application/pdf",
                            filename=file_name)
    else:
        return JSONResponse({"message": "File not found"}, status_code=404)
